from dotenv import load_dotenv
load_dotenv()
from flask import (
    Flask, render_template, request, redirect, session, url_for, flash, jsonify
)
import mysql.connector
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
import os
from datetime import datetime
import json
import re
import requests
from bs4 import BeautifulSoup
import subprocess
import mediapipe as mp
import cv2
# =============================
# External Service Config
# =============================
ASSEMBLYAI_API_KEY = os.getenv("ASSEMBLYAI_API_KEY")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
# Optional imports with fallbacks
try:
    import nltk
    from nltk.sentiment import SentimentIntensityAnalyzer
    VADER_READY = True
    try:
        _ = SentimentIntensityAnalyzer()
    except Exception:
        nltk.download('vader_lexicon')
        _ = SentimentIntensityAnalyzer()
except Exception as e:
    print(f"NLTK/VADER not available: {e}")
    VADER_READY = False

try:
    from openai import OpenAI
    OPENAI_AVAILABLE = True
    client = OpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None
except ImportError:
    OPENAI_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

app = Flask(__name__)
app.secret_key = os.getenv('SECRET_KEY', 'interviewwise-secret-key')

# Upload directories
UPLOAD_FOLDER = 'static/uploads'
AUDIO_UPLOAD_DIR = os.path.join('static', 'interview_audio')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(AUDIO_UPLOAD_DIR, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
FFMPEG_PATH="C:/Users/admin/Downloads/ffmpeg-8.0-essentials_build/ffmpeg-8.0-essentials_build/bin/ffmpeg.exe"

# =============================
# Database Connection
# =============================
def get_db_connection():
    try:
        return mysql.connector.connect(
            host="localhost",
            user="root",
            password="1234",
            database="interviewwise_db"
        )
    except mysql.connector.Error as err:
        print(f"DB Error: {err}")
        return None

# =============================
# MediaPipe Setup
# =============================
mp_face = mp.solutions.face_detection.FaceDetection(min_detection_confidence=0.5)
mp_pose = mp.solutions.pose.Pose(min_detection_confidence=0.5)

def analyze_frame(frame):
    """Run MediaPipe analysis on a frame."""
    face_results = mp_face.process(cv2.cvtColor(frame, cv2.COLOR_BGR2RGB))
    pose_results = mp_pose.process(cv2.cvtColor(frame, cv2.COLOR_BGR2RGB))

    face_count = len(face_results.detections) if face_results.detections else 0
    pose_detected = bool(pose_results.pose_landmarks)

    return face_count, pose_detected

def analyze_video_metrics(video_path):
    """
    Analyze a video for face and body metrics using MediaPipe.
    Returns a dictionary with all relevant metrics.
    """
    metrics = {
        "eye_contact_percent": 0.0,
        "smile_frequency": 0.0,
        "posture_score": 0.0,
        "head_tilt_score": 0.0,
        "blink_rate": 0.0,
        "speaking_pace": 0.0,
        "voice_stability": 0.0,
        "overall_body_language_score": 0.0,
        "attention_span_score": 0.0,
        "filler_word_count": 0
    }

    try:
        cap = cv2.VideoCapture(video_path)
        if not cap.isOpened():
            print(f"❌ Cannot open video: {video_path}")
            return metrics

        face_detector = mp.solutions.face_detection.FaceDetection(min_detection_confidence=0.5)
        pose_detector = mp.solutions.pose.Pose(min_detection_confidence=0.5)

        total_frames = 0
        faces_detected = 0
        smiles_detected = 0
        posture_sum = 0.0
        head_tilt_sum = 0.0
        blinks_detected = 0

        while True:
            ret, frame = cap.read()
            if not ret:
                break
            total_frames += 1
            rgb_frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)

            # Face
            face_results = face_detector.process(rgb_frame)
            if face_results.detections:
                faces_detected += 1
                smiles_detected += 1  # placeholder
                head_tilt_sum += 0.0
                blinks_detected += 0  # placeholder

            # Pose
            pose_results = pose_detector.process(rgb_frame)
            if pose_results.pose_landmarks:
                posture_sum += 1.0

        cap.release()
        face_detector.close()
        pose_detector.close()

        if total_frames > 0:
            metrics["eye_contact_percent"] = round(faces_detected / total_frames * 100, 2)
            metrics["smile_frequency"] = round(smiles_detected / total_frames * 100, 2)
            metrics["posture_score"] = round(posture_sum / total_frames * 100, 2)
            metrics["head_tilt_score"] = round(head_tilt_sum / total_frames, 2)
            metrics["blink_rate"] = round(blinks_detected / total_frames, 2)
            metrics["overall_body_language_score"] = round(
                (metrics["eye_contact_percent"] + metrics["smile_frequency"] + metrics["posture_score"])/3,2
            )

        print("✅ Video metrics computed:", metrics)
        return metrics

    except Exception as e:
        print("❌ Error in analyze_video_metrics:", e)
        return metrics


# =============================
# Cache Control
# =============================
@app.after_request
def add_header(response):
    response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, public, max-age=0'
    response.headers['Pragma'] = 'no-cache'
    response.headers['Expires'] = '0'
    return response

# =============================
# AssemblyAI Transcription
# =============================
def transcribe_with_assemblyai(file_path: str) -> str:
    """
    Upload audio file to AssemblyAI and return transcript text.
    Logs each step so you know exactly where it fails.
    """
    if not ASSEMBLYAI_API_KEY:
        print("❌ AssemblyAI key missing.")
        return "(AssemblyAI key missing)"

    headers = {"authorization": ASSEMBLYAI_API_KEY}
    try:
        # Step 1: Upload file
        print(f"📤 Uploading {file_path} to AssemblyAI...")
        with open(file_path, "rb") as f:
            upload_res = requests.post(
                "https://api.assemblyai.com/v2/upload",
                headers=headers,
                data=f
            )
        if upload_res.status_code != 200:
            print("❌ Upload failed:", upload_res.text)
            return "(upload failed)"
        upload_url = upload_res.json().get("upload_url")
        print("✅ Uploaded. URL:", upload_url)

        # Step 2: Request transcription
        transcript_res = requests.post(
            "https://api.assemblyai.com/v2/transcript",   # ✅ correct endpoint
            headers=headers,
            json={"audio_url": upload_url}
        )
        if transcript_res.status_code != 200:
            print("❌ Transcription request failed:", transcript_res.text)
            return "(transcription request failed)"

        transcript_id = transcript_res.json().get("id")
        if not transcript_id:
            print("❌ No transcript ID in response:", transcript_res.json())
            return "(no transcript id)"
        print("🆔 Transcription job ID:", transcript_id)

        # Step 3: Poll until complete
        while True:
            poll_res = requests.get(
                f"https://api.assemblyai.com/v2/transcript/{transcript_id}",  # ✅ corrected
                headers=headers
            )
            if poll_res.status_code != 200:
                print("❌ Polling failed:", poll_res.text)
                return "(polling failed)"

            data = poll_res.json()
            status = data.get("status")
            if status == "completed":
                print("✅ Transcription completed.")
                return data.get("text", "").strip()
            elif status == "error":
                print("❌ Transcription error:", data)
                return "(transcription error)"

            # keep polling every 3s
            import time
            time.sleep(3)

    except Exception as e:
        print("[AssemblyAI exception]", e)
        return "(transcription unavailable)"


# =============================
# Analysis Helpers
# =============================
FILLER_WORDS = {
    'um','uh','like','you know','so','actually','basically','literally',
    'kinda','sorta','right','okay','ok','well'
}

def count_filler_words(text: str) -> int:
    if not text:
        return 0
    t = text.lower()
    count = 0
    for fw in ['you know']:
        count += len(re.findall(r'\b' + re.escape(fw) + r'\b', t))
        t = re.sub(r'\b' + re.escape(fw) + r'\b', ' ', t)
    for w in FILLER_WORDS - {'you know'}:
        count += len(re.findall(r'\b' + re.escape(w) + r'\b', t))
    return count

def analyze_sentiment(text: str):
    if not text:
        return ("Neutral", 0.00)
    if VADER_READY:
        sia = SentimentIntensityAnalyzer()
        scores = sia.polarity_scores(text)
        compound = scores.get('compound', 0.0)
        label = "Positive" if compound >= 0.05 else "Negative" if compound <= -0.05 else "Neutral"
        return (label, round(abs(compound), 2))
    return ("Neutral", 0.00)

def simple_confidence_proxy(text: str, duration_sec: float, filler_count: int) -> float:
    words = len(text.split()) if text else 0
    raw = (min(words / 120.0, 1.0) * 0.5) + (min(duration_sec / 60.0, 1.0) * 0.3) + (max(1 - (filler_count/15.0), 0) * 0.2)
    return round(max(0.0, min(raw, 1.0)), 2)

# =============================
# LLM Feedback (OpenAI)
# =============================
def analyze_answer_with_llm(question_text: str, user_answer: str, user_role: str = None):
    if not OPENAI_AVAILABLE:
        return {"score": None, "feedback": "LLM unavailable.", "suggestions": []}
    if not OPENAI_API_KEY:
        return {"score": None, "feedback": "OPENAI_API_KEY not configured.", "suggestions": []}
    try:
        client = OpenAI(api_key=OPENAI_API_KEY)
        role_context = f" for a {user_role} position" if user_role else ""
        prompt = (
            f"You are an expert interview evaluator{role_context}. "
            f"Analyze this interview response and return JSON with: "
            f"score (0-10), feedback (constructive critique), "
            f"suggestions (array of 3 specific improvement tips), "
            f"technical_accuracy (0-10 if applicable), "
            f"communication_clarity (0-10).\n\n"
            f"Question: {question_text}\n"
            f"Answer: {user_answer}\n"
        )
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"}
        )
        data = json.loads(resp.choices[0].message.content)
        if "score" not in data or "feedback" not in data:
            raise ValueError("LLM returned unexpected format")
        return data
    except Exception as e:
        print(f"[LLM error] {e}")
        return {"score": None, "feedback": "Could not generate detailed AI feedback.", "suggestions": []}

def generate_quiz_questions(parsed_resume, role):
    if not OPENAI_AVAILABLE:
        return []

    try:
        client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
        skills_str = ', '.join(parsed_resume.get('skills', []))
        experience = parsed_resume.get('experience_years', 0)

        prompt = f"""
        Generate 3 multiple-choice quiz questions (MCQs) for a {role} candidate with:
        - Skills: {skills_str}
        - Experience: {experience} years

        Each question must include:
        - question (string)
        - option_a, option_b, option_c, option_d
        - correct_option (one of A/B/C/D)

        Return ONLY valid JSON.
        """

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"}
        )

        raw = response.choices[0].message.content
        print("[AI raw quiz output]", raw)   # ✅ Debug log

        data = json.loads(raw)

        if isinstance(data, list):
            return data
        elif "questions" in data:
            return data["questions"]
        else:
            return []

    except Exception as e:
        print(f"Error generating quiz questions: {e}")
        return []

# -----------------------------
# Resume Parsing Functions
# -----------------------------
class ResumeParser:
    def __init__(self):
        self.common_skills = {
            'programming': ['python', 'java', 'javascript', 'c++', 'c#', 'php', 'ruby', 'go', 'rust', 'swift'],
            'web': ['html', 'css', 'react', 'angular', 'vue', 'node.js', 'django', 'flask', 'express'],
            'data': ['sql', 'mysql', 'postgresql', 'mongodb', 'pandas', 'numpy', 'matplotlib', 'tableau', 'power bi'],
            'cloud': ['aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'terraform'],
            'ml': ['machine learning', 'deep learning', 'tensorflow', 'pytorch', 'scikit-learn', 'opencv']
        }
    
    def extract_skills_from_text(self, text):
        text_lower = text.lower()
        found_skills = []
        
        for category, skills in self.common_skills.items():
            for skill in skills:
                if skill.lower() in text_lower:
                    found_skills.append(skill)
        
        return list(set(found_skills))
    
    def extract_experience_years(self, text):
        experience_patterns = [
            r'(\\d+)\\+?\\s*years?\\s*of\\s*experience',
            r'(\\d+)\\+?\\s*years?\\s*experience',
            r'experience\\s*of\\s*(\\d+)\\+?\\s*years?',
        ]
        
        years = []
        for pattern in experience_patterns:
            matches = re.findall(pattern, text.lower())
            years.extend([int(match) for match in matches])
        
        return max(years) if years else 0
    
    def parse_resume_file(self, file_path):
        try:
            if file_path.endswith('.pdf') and PDF_AVAILABLE:
                text = self._extract_text_from_pdf(file_path)
            elif file_path.endswith(('.doc', '.docx')) and DOCX_AVAILABLE:
                text = self._extract_text_from_docx(file_path)
            else:
                return None
            
            skills = self.extract_skills_from_text(text)
            experience_years = self.extract_experience_years(text)
            
            return {
                'skills': skills,
                'experience_years': experience_years,
                'raw_text': text[:1000]
            }
            
        except Exception as e:
            print(f"Error parsing resume: {e}")
            return None
        
    def scrape_linkedin(self, url):
        """⚠️ Basic scraper (works only on public LinkedIn profiles).
           For private profiles, need LinkedIn API/selenium."""
        try:
            headers = {"User-Agent": "Mozilla/5.0"}
            resp = requests.get(url, headers=headers)
            if resp.status_code != 200:
                return {"skills": [], "summary": ""}
            soup = BeautifulSoup(resp.text, 'html.parser')
            text = soup.get_text(" ", strip=True)
            return {
                "skills": self.extract_skills_from_text(text),
                "summary": text[:2000]
            }
        except Exception as e:
            print(f"LinkedIn scrape error: {e}")
            return {"skills": [], "summary": ""}
        
    def _extract_text_from_pdf(self, file_path):
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ''
                for page in reader.pages:
                    text += page.extract_text()
                return text
        except Exception as e:
            print(f"Error extracting PDF text: {e}")
            return ""
    
    def _extract_text_from_docx(self, file_path):
        try:
            doc = Document(file_path)
            text = ''
            for paragraph in doc.paragraphs:
                text += paragraph.text + '\\n'
            return text
        except Exception as e:
            print(f"Error extracting DOCX text: {e}")
            return ""

def generate_resume_based_questions(parsed_resume, role):
    if not OPENAI_AVAILABLE:
        return []

    try:
        client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
        skills_str = ', '.join(parsed_resume.get('skills', []))
        experience = parsed_resume.get('experience_years', 0)

        prompt = f"""
        Generate 3 interview questions for a {role} candidate with:
        - Skills: {skills_str}
        - Experience: {experience} years

        Return ONLY valid JSON. Each item must have:
        - question_text
        - difficulty (Easy/Medium/Hard)
        - category (Technical/Behavioral)
        """

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}]
        )

        raw = response.choices[0].message.content
        print("==== RAW LLM OUTPUT (resume) ====")
        print(raw)
        print("=================================")

        try:
            data = json.loads(raw)
        except Exception as e:
            match = re.search(r"\[.*\]", raw, re.S)
            if match:
                try:
                    data = json.loads(match.group(0))
                except Exception as e2:
                    print("JSON parse error:", e2)
                    return []
            else:
                print("Could not parse AI output into JSON")
                return []

        return data.get('questions', []) if isinstance(data, dict) else data

        # Handle both wrapped & direct array
        if isinstance(data, list):
            return data
        elif "questions" in data:
            return data["questions"]
        else:
            return []

    except Exception as e:
        print(f"Error generating interview questions: {e}")
        return []



def calculate_enhanced_score(user_answer, question_text, user_role, metrics, ai_feedback):
    conn = get_db_connection()
    if not conn:
        return ai_feedback.get('score', 5)
    
    try:
        cur = conn.cursor(dictionary=True)
        cur.execute("SELECT * FROM scoring_criteria WHERE role = %s", (user_role,))
        criteria = cur.fetchall()
        
        if not criteria:
            return ai_feedback.get('score', 5)
        
        # Calculate weighted score based on role-specific criteria
        total_score = 0
        total_weight = 0
        
        for criterion in criteria:
            weight = float(criterion['weight'])
            if criterion['criteria_name'] == 'Technical Knowledge':
                score = ai_feedback.get('technical_accuracy', 5)
            elif criterion['criteria_name'] == 'Communication':
                score = ai_feedback.get('communication_clarity', 5)
            elif 'Body Language' in criterion['criteria_name']:
                score = (metrics.get('eye_contact_percent', 0) / 10 + 
                        metrics.get('posture_score', 0) / 10) / 2
            else:
                score = ai_feedback.get('score', 5)
            
            total_score += score * weight
            total_weight += weight
        
        return round(total_score / total_weight, 1) if total_weight > 0 else 5
        
    except Exception as e:
        print(f"Error calculating enhanced score: {e}")
        return ai_feedback.get('score', 5)
    finally:
        if conn:
            conn.close()

def extract_text_from_file(file_path):
    text = ""
    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext == ".pdf":
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() + "\n"

        elif ext in [".doc", ".docx"]:
            doc = doc.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"

        else:
            text = "Unsupported file format for AI review."

    except Exception as e:
        text = f"Error extracting text: {e}"

    return text.strip()

# -----------------------------
# Main Routes
# -----------------------------
@app.route('/')
def index():
    user_name = session.get('user')
    return render_template('index.html', user=user_name)

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email_or_username = request.form.get('username')
        password = request.form.get('password')

        conn = get_db_connection()
        if not conn:
            flash("Database connection error!", "danger")
            return redirect(url_for('login'))

        cur = conn.cursor(dictionary=True)
        cur.execute("""
            SELECT * FROM users
            WHERE email=%s OR username=%s
        """, (email_or_username, email_or_username))
        user = cur.fetchone()
        cur.close()
        conn.close()

        if user and check_password_hash(user['password'], password):
            session['user'] = user['username']
            session['user_id'] = user['id']
            session['role'] = user['role'] or None

            flash("Login successful!", "success")
            return redirect(url_for('index'))
        else:
            flash("Invalid credentials!", "danger")
    return render_template('login.html')

# -------------------------
# Register
# -------------------------
@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        email = request.form.get('email', '').strip()
        password = request.form.get('password', '')
        confirm_password = request.form.get('confirm_password', '')
        role = request.form.get('role')
        education_level = request.form.get('graduation')
        linkedin_profile = request.form.get('linkedin')
        resume_file = request.files.get('resume')

        if not username or not email or not password:
            flash("Please fill required fields.", "danger")
            return redirect(url_for('register'))

        if password != confirm_password:
            flash("Passwords do not match!", "danger")
            return redirect(url_for('register'))

        hashed_password = generate_password_hash(password)

        resume_path = None
        if resume_file and resume_file.filename:
            ts = datetime.now().strftime("%Y%m%d%H%M%S")
            safe_name = f"{secure_filename(username)}_{ts}_{secure_filename(resume_file.filename)}"
            target = os.path.join(app.config['UPLOAD_FOLDER'], safe_name)
            try:
                resume_file.save(target)
                # store relative path (to static/) for easier serving
                resume_path = f"uploads/{safe_name}"
            except Exception as e:
                print("Resume save error:", e)
                resume_path = None

        conn = get_db_connection()
        if not conn:
            flash("Database connection error!", "danger")
            return redirect(url_for('register'))

        cur = conn.cursor(dictionary=True)
        try:
            cur.execute("""
                INSERT INTO users 
                (username, email, password, role, education_level, linkedin_profile, resume_path, registered_at)
                VALUES (%s, %s, %s, %s, %s, %s, %s, NOW())
            """, (username, email, hashed_password, role, education_level, linkedin_profile, resume_path))
            conn.commit()

            cur.execute("SELECT * FROM users WHERE username=%s", (username,))
            new_user = cur.fetchone()

            session['user'] = new_user['username']
            session['user_id'] = new_user['id']
            session['role'] = new_user.get('role')
            flash("Registration successful! Welcome, " + username, "success")
            return redirect(url_for('index'))

        except mysql.connector.IntegrityError as e:
            print("Register DB error:", e)
            flash("Email or username may already exist.", "danger")
            # If we saved a resume but DB insert failed — try to remove the file
            if resume_path:
                try:
                    os.remove(os.path.join(app.root_path, 'static', resume_path))
                except Exception:
                    pass
        finally:
            cur.close()
            conn.close()

    return render_template('register.html')


@app.route('/profile')
def profile():
    if 'user' not in session:
        return redirect(url_for('login'))

    conn = get_db_connection()
    if not conn:
        flash("Database connection error!", "danger")
        return redirect(url_for('login'))

    cur = conn.cursor(dictionary=True)

    # Fetch user data
    cur.execute("SELECT * FROM users WHERE username=%s", (session['user'],))
    user_data = cur.fetchone()

    # Quiz progress
    cur.execute("""
        SELECT DATE(attempted_at) AS date,
               SUM(is_correct) AS correct,
               COUNT(*) AS total
        FROM quiz_attempts
        WHERE user_id = %s
        GROUP BY DATE(attempted_at)
        ORDER BY date
    """, (user_data['id'],))
    progress_data = cur.fetchall()

    # Fetch interview attempts
    cur.execute("""
        SELECT lia.id,
               lia.attempted_at,
               COALESCE(iq.question_text, dq.generated_question) AS question,
               lia.score_overall,
               lia.eye_contact_percent,
               lia.smile_frequency,
               lia.attention_span_score,
               lia.blink_rate,
               lia.posture_score,
               lia.head_tilt_score
        FROM live_interview_attempts lia
        LEFT JOIN interview_questions iq ON lia.question_id = iq.id
        LEFT JOIN dynamic_questions dq ON lia.question_id = dq.id
        WHERE lia.user_id = %s
        ORDER BY lia.attempted_at DESC
        LIMIT 10
    """, (user_data['id'],))
    attempts = cur.fetchall()

    # Fetch feedback separately
    interview_history = []
    for att in attempts:
        cur.execute("""
            SELECT feedback_category, feedback_text, score
            FROM interview_feedback_detailed
            WHERE interview_attempt_id = %s
        """, (att['id'],))
        feedback_rows = cur.fetchall()

        # Organize feedback
        feedback = {
            "Technical": None,
            "Communication": None,
            "Body Language": None,
            "Overall": None
        }
        for row in feedback_rows:
            if row['feedback_category'] in feedback:
                if row['feedback_category'] == 'Overall':
                    feedback[row['feedback_category']] = row['feedback_text']  # full text
                else:
                    feedback[row['feedback_category']] = row['score']  # only numeric score

        att['technical_feedback'] = feedback['Technical']
        att['communication_feedback'] = feedback['Communication']
        att['body_language_feedback'] = feedback['Body Language']
        att['overall_feedback'] = feedback['Overall']

        interview_history.append(att)

    # Resume analysis
    cur.execute("SELECT * FROM resume_analysis WHERE user_id = %s", (user_data['id'],))
    resume_analysis = cur.fetchone()

    cur.close()
    conn.close()

    return render_template(
        'profile.html',
        user=user_data,
        progress_data=progress_data,
        live_interview_history=interview_history,
        resume_analysis=resume_analysis
    )

# -------------------------
# Edit Profile (full robust)
# -------------------------
@app.route('/edit_profile', methods=['GET', 'POST'])
def edit_profile():
    if 'user' not in session:
        return redirect(url_for('login'))

    conn = get_db_connection()
    if not conn:
        flash("Database connection error!", "danger")
        return redirect(url_for('profile'))

    cur = conn.cursor(dictionary=True)
    cur.execute("SELECT * FROM users WHERE username=%s", (session['user'],))
    user_data = cur.fetchone()

    if request.method == 'POST':
        new_email = request.form.get('email') or user_data['email']
        new_role = request.form.get('role') or user_data['role']
        new_education = request.form.get('education_level') or user_data['education_level']
        new_linkedin = request.form.get('linkedin_profile') or user_data['linkedin_profile']
        resume_file = request.files.get('resume')

        resume_path = user_data.get('resume_path')  # default keep

        # handle new resume upload
        if resume_file and resume_file.filename:
            # delete old resume if exists and is stored under static/uploads
            if user_data.get('resume_path'):
                old_full = os.path.join(app.root_path, 'static', user_data['resume_path'])
                if os.path.exists(old_full):
                    try:
                        os.remove(old_full)
                    except Exception as e:
                        print("Could not delete old resume:", e)

            ts = datetime.now().strftime("%Y%m%d%H%M%S")
            filename = f"{secure_filename(session['user'])}_{ts}_{secure_filename(resume_file.filename)}"
            full_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            try:
                resume_file.save(full_path)
                resume_path = f"uploads/{filename}"
            except Exception as e:
                print("Error saving new resume:", e)
                flash("Could not save resume file.", "warning")

        if not new_email:
            flash("Email is required!", "danger")
            cur.close(); conn.close()
            return redirect(url_for('edit_profile'))

        cur.execute("""
            UPDATE users
            SET email=%s, role=%s, education_level=%s, linkedin_profile=%s, resume_path=%s
            WHERE id=%s
        """, (new_email, new_role, new_education, new_linkedin, resume_path, user_data['id']))
        conn.commit()

        cur.close()
        conn.close()
        session['role'] = new_role
        flash("Profile updated successfully!", "success")
        return redirect(url_for('profile'))

    cur.close()
    conn.close()
    return render_template('edit_profile.html', user=user_data)

# -------------------------
# Review Uploaded Resume + Cover Letter with AI
# -------------------------
@app.route("/review_documents")
def review_documents():
    if "user_id" not in session:
        flash("Please login first.")
        return redirect(url_for("login"))

    # Open DB connection
    conn = get_db_connection()
    if conn is None:
        flash("Database connection failed.")
        return redirect(url_for("profile"))

    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT * FROM users WHERE id = %s", (session["user_id"],))
    user = cursor.fetchone()
    cursor.close()
    conn.close()

    if not user:
        flash("User not found.")
        return redirect(url_for("login"))

    resume_text = ""
    cover_text = ""

    # Extract text from resume if exists
    if user["resume_path"]:
        resume_file = os.path.join("static", user["resume_path"])
        resume_text = extract_text_from_file(resume_file)

    # Extract text from cover letter if exists
    if user.get("cover_letter_path"):
        cover_file = os.path.join("static", user["cover_letter_path"])
        cover_text = extract_text_from_file(cover_file)

    if not resume_text and not cover_text:
        flash("No resume or cover letter uploaded.")
        return redirect(url_for("profile"))

    # Build AI prompt
    prompt = f"""
    You are a professional career coach. Review the following candidate documents:

    Resume:
    {resume_text[:4000]}

    Cover Letter:
    {cover_text[:4000]}

    Provide a detailed review highlighting strengths, weaknesses, and job-readiness.
    """

    ai_review = ""
    try:
        # ✅ FIX: use `prompt` instead of undefined `user_answer`
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an AI feedback assistant."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=600
        )
        ai_review = response.choices[0].message.content.strip()

    except Exception as e:
        ai_review = f"Error fetching AI review: {e}"

    # Save AI review into resume_analysis table
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO resume_analysis (user_id, resume_text, cover_letter_text, ai_review)
            VALUES (%s, %s, %s, %s)
            ON DUPLICATE KEY UPDATE
                resume_text=%s,
                cover_letter_text=%s,
                ai_review=%s,
                updated_at=NOW()
        """, (
            session["user_id"], resume_text, cover_text, ai_review,
            resume_text, cover_text, ai_review
        ))
        conn.commit()
        cur.close()
        conn.close()
    except Exception as e:
        print("Error saving AI review to DB:", e)

    return render_template("review.html", user=user, ai_review=ai_review)


# -------------------------
# Analyze Resume + LinkedIn, generate questions, safe JSON extraction
# -------------------------
@app.route('/analyze_resume')
def analyze_resume():
    if 'user' not in session:
        return redirect(url_for('login'))

    user_id = session.get('user_id')
    conn = get_db_connection()
    if not conn:
        flash("Database connection error!", "danger")
        return redirect(url_for('profile'))

    cur = conn.cursor(dictionary=True)
    cur.execute("SELECT id, username, email, resume_path, linkedin_profile, role FROM users WHERE id = %s", (user_id,))
    user = cur.fetchone()

    if not user or not user.get('resume_path'):
        flash("No resume found. Please upload a resume first.", "warning")
        cur.close(); conn.close()
        return redirect(url_for('edit_profile'))

    resume_parser = ResumeParser()
    full_resume_path = os.path.join(app.root_path, 'static', user['resume_path'])

    if not os.path.exists(full_resume_path):
        flash("Resume file not found on server.", "danger")
        cur.close(); conn.close()
        return redirect(url_for('profile'))

    # Parse resume
    analysis = resume_parser.parse_resume_file(full_resume_path)
    if not analysis or not isinstance(analysis, dict):
        flash("Failed to parse resume file.", "danger")
        cur.close(); conn.close()
        return redirect(url_for('profile'))

    # LinkedIn scraping (optional)
    if user.get('linkedin_profile'):
        try:
            linkedin_analysis = resume_parser.scrape_linkedin(user['linkedin_profile'])
            if linkedin_analysis:
                combined = set(analysis.get('skills', [])) | set(linkedin_analysis.get('skills', []))
                analysis['skills'] = list(combined)
                analysis['linkedin_summary'] = linkedin_analysis.get('summary', '')
        except Exception as e:
            print("LinkedIn scrape error:", e)

    # AI Review
    try:
        skills = ", ".join(analysis.get('skills', []))
        exp_years = analysis.get('experience_years', 0)
        linkedin_summary = analysis.get('linkedin_summary', '')
        role = user.get('role', '')

        prompt = f"""
        You are an AI career coach. Analyze the following candidate resume:

        Role applied: {role}
        Skills: {skills}
        Years of Experience: {exp_years}
        LinkedIn Summary: {linkedin_summary}

        Provide a structured review with:
        1. Strengths
        2. Weaknesses
        3. Missing skills (important for the role)
        4. Suggestions for improvement
        5. Overall rating: Poor, Moderate, Good, or Excellent
        """

        response = client.chat.completions.create(
            model="gpt-4o-mini",   # ✅ correct spelling (letter o, not zero)
            messages=[
                {"role": "system", "content": "You are an AI feedback assistant."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=600
        )


        ai_review = response.choices[0].message.content.strip()
        analysis['ai_review'] = ai_review

    except Exception as e:
        print("AI review generation error:", e)
        analysis['ai_review'] = "AI review not available."

    # Save to DB
    try:
        cur.execute("""
            INSERT INTO resume_analysis (user_id, parsed_skills, experience_years, key_technologies, linkedin_summary, ai_review)
            VALUES (%s, %s, %s, %s, %s, %s)
            ON DUPLICATE KEY UPDATE
                parsed_skills=%s,
                experience_years=%s,
                key_technologies=%s,
                linkedin_summary=%s,
                ai_review=%s,
                updated_at=NOW()
        """, (
            user_id,
            json.dumps(analysis.get('skills', [])),
            analysis.get('experience_years', 0),
            json.dumps(analysis.get('skills', [])),
            analysis.get('linkedin_summary', ''),
            analysis.get('ai_review', ''),

            json.dumps(analysis.get('skills', [])),
            analysis.get('experience_years', 0),
            json.dumps(analysis.get('skills', [])),
            analysis.get('linkedin_summary', ''),
            analysis.get('ai_review', '')
        ))

        conn.commit()
        flash("Resume analyzed successfully with AI review.", "success")

    except Exception as e:
        print("Error saving resume_analysis:", e)
        flash("Failed to save resume analysis.", "danger")

    finally:
        cur.close()
        conn.close()

    # Render review page
    return render_template("resume_review.html", ai_review=analysis['ai_review'], user=user)

    # ⚠️ NOTE: your dynamic questions + quiz code is unchanged


@app.route('/quiz', methods=['GET','POST'])
def quiz():
    if 'user' not in session:
        return redirect(url_for('login'))

    conn = get_db_connection()
    if not conn:
        flash("Database connection error!", "danger")
        return redirect(url_for('login'))

    cur = conn.cursor(dictionary=True)
    cur.execute("SELECT * FROM users WHERE username=%s", (session['user'],))
    user = cur.fetchone()

    cur.execute("SELECT * FROM quizzes WHERE role=%s ORDER BY RAND() LIMIT 5", (user['role'],))
    questions = cur.fetchall()

    if request.method == 'POST':
        score = 0
        total = 0
        for key, value in request.form.items():
            if key.startswith('question_'):
                qid = int(key.split('_')[1])
                selected = value
                cur.execute("SELECT correct_option FROM quizzes WHERE id=%s", (qid,))
                row = cur.fetchone()
                if row:
                    is_correct = (selected == row['correct_option'])
                    score += int(is_correct)
                    total += 1
                    cur.execute("""
                        INSERT INTO quiz_attempts (user_id, quiz_id, selected_option, is_correct, attempted_at)
                        VALUES (%s,%s,%s,%s,NOW())
                    """, (user['id'], qid, selected, is_correct))
        conn.commit()
        cur.close()
        conn.close()
        return render_template('quiz_result.html', score=score, total=total)

    cur.close()
    conn.close()
    return render_template('quiz.html', role=user['role'], questions=questions)

@app.route("/mock_interview")
def mock_interview():
    if "user" not in session:
        return redirect(url_for("login"))

    user_role = session.get("role")
    if not user_role:
        flash("No role found. Please set your role.", "warning")
        return redirect(url_for("profile"))

    conn = get_db_connection()
    cur = conn.cursor(dictionary=True)
    user_id = session.get("user_id")

    # Personalized dynamic question
    cur.execute(
        "SELECT id, generated_question AS question_text FROM dynamic_questions WHERE user_id=%s ORDER BY RAND() LIMIT 1",
        (user_id,)
    )
    root = cur.fetchone()
    is_personalized = bool(root)

    if not root:
        cur.execute(
            "SELECT id, question_text FROM interview_questions WHERE role=%s ORDER BY RAND() LIMIT 1",
            (user_role,)
        )
        root = cur.fetchone()

    if not root:
        cur.close(); conn.close()
        flash("No questions available.", "warning")
        return redirect(url_for("profile"))

    # Prepare follow-up queue
    followups = []
    if is_personalized:
        cur.execute(
            "SELECT id, generated_question AS question_text FROM dynamic_questions WHERE user_id=%s AND id!=%s ORDER BY RAND() LIMIT 2",
            (user_id, root['id'])
        )
        followups = cur.fetchall()

    remaining = 3 - len(followups)
    if remaining > 0:
        cur.execute(
            "SELECT id, question_text FROM interview_questions WHERE role=%s ORDER BY RAND() LIMIT %s",
            (user_role, remaining)
        )
        followups.extend(cur.fetchall())

    session["mi_root"] = root["id"]
    session["mi_queue"] = [f["id"] for f in followups]

    cur.close(); conn.close()

    return render_template(
        "mock_interview.html",
        root_question=root,
        total_questions=1 + len(followups),
        is_personalized=is_personalized
    )

# =============================
# Upload Interview Response
# =============================
@app.route("/api/live_interview/upload", methods=["POST"])
def upload_interview_response():
    try:
        user_id = session.get("user_id")
        if not user_id:
            return jsonify({"error": "Not logged in"}), 401

        question_id = request.form.get("question_id")
        if not question_id:
            return jsonify({"error": "Missing question_id"}), 400

        file = request.files.get("audio")
        if not file:
            return jsonify({"error": "No audio file uploaded"}), 400

        os.makedirs(AUDIO_UPLOAD_DIR, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        raw_filename = f"user{user_id}_q{question_id}_{timestamp}.webm"
        raw_path = os.path.join(AUDIO_UPLOAD_DIR, raw_filename)
        file.save(raw_path)
        print("📂 Saved audio/video to:", raw_path)

        # Convert to WAV
        wav_filename = raw_filename.replace(".webm", ".wav")
        wav_path = os.path.join(AUDIO_UPLOAD_DIR, wav_filename)
        ffmpeg_cmd = [
            FFMPEG_PATH, "-i", raw_path, "-vn", "-acodec", "pcm_s16le", "-ar", "48000", "-ac", "1", wav_path
        ]
        try:
            subprocess.run(ffmpeg_cmd, check=True)
            print("🎶 Converted to WAV:", wav_path)
        except Exception as e:
            print("❌ FFmpeg conversion failed:", e)
            wav_path = None

        # Transcribe audio
        transcript_text = transcribe_with_assemblyai(wav_path) if wav_path else ""
        filler_count = count_filler_words(transcript_text)
        duration_sec = request.form.get("duration_sec", 0)

        # Video metrics using MediaPipe
        metrics = analyze_video_metrics(raw_path)

        # Fetch question text
        conn = get_db_connection()
        cur = conn.cursor(dictionary=True)
        cur.execute("SELECT question_text FROM interview_questions WHERE id=%s", (question_id,))
        qrow = cur.fetchone()
        question_text = qrow["question_text"] if qrow else "Unknown question"

        # LLM feedback
        llm_result = analyze_answer_with_llm(question_text, transcript_text, session.get("role"))

        ai_feedback_text = llm_result.get("feedback", "")
        overall_score = llm_result.get("score", 0)
        technical_score = llm_result.get("technical_accuracy", 0)
        communication_score = llm_result.get("communication_clarity", 0)
        body_language_score = llm_result.get("body_language", 0)
        suggestions_json = json.dumps(llm_result.get("suggestions", []))

        # Insert into live_interview_attempts
        sql = """
            INSERT INTO live_interview_attempts
            (user_id, question_id, user_answer_text, audio_path, video_path,
             ai_feedback, score_overall, sentiment_label, sentiment_score, confidence_score,
             filler_word_count, eye_contact_percent, smile_frequency, head_tilt_score, blink_rate,
             posture_score, speaking_pace, voice_stability, overall_body_language_score,
             attention_span_score, attempted_at, duration_sec)
            VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,NOW(),%s)
        """
        cur.execute(sql, (
            user_id, question_id, transcript_text, raw_path, raw_path,
            ai_feedback_text, overall_score, None, None, None,
            filler_count,
            metrics.get("eye_contact_percent", 0), metrics.get("smile_frequency", 0),
            metrics.get("head_tilt_score", 0), metrics.get("blink_rate", 0),
            metrics.get("posture_score", 0), metrics.get("speaking_pace", 0),
            metrics.get("voice_stability", 0), metrics.get("overall_body_language_score", 0),
            metrics.get("attention_span_score", 0), duration_sec
        ))
        conn.commit()
        attempt_id = cur.lastrowid

        # Insert feedback for all categories
        feedback_categories = [
            ("Technical", f"Score: {technical_score}", technical_score),
            ("Communication", f"Score: {communication_score}", communication_score),
            ("Body Language", f"Score: {body_language_score}", body_language_score),
            ("Overall", ai_feedback_text, overall_score)
        ]

        for category, text, score in feedback_categories:
            cur.execute("""
                INSERT INTO interview_feedback_detailed
                (interview_attempt_id, feedback_category, feedback_text, improvement_suggestions, score)
                VALUES (%s,%s,%s,%s,%s)
            """, (attempt_id, category, text, suggestions_json, score))
        conn.commit()

        cur.close()
        conn.close()

        return jsonify({
            "message": "Interview attempt saved with AI feedback!",
            "transcript": transcript_text,
            "ai_feedback": ai_feedback_text,
            "overall_score": overall_score,
            "technical_score": technical_score,
            "communication_score": communication_score,
            "body_language_score": body_language_score,
            "metrics": metrics,
            "suggestions": json.loads(suggestions_json)
        })

    except Exception as e:
        print("❌ Error in /api/live_interview/upload:", e)
        return jsonify({"error": "Failed to save interview", "details": str(e)}), 500

# =============================
# Get Question Text API
# =============================
@app.route("/api/get_question_text/<int:qid>")
def get_question_text(qid):
    conn = get_db_connection()
    cur = conn.cursor(dictionary=True)
    cur.execute("SELECT question_text FROM interview_questions WHERE id=%s", (qid,))
    row = cur.fetchone()
    cur.close(); conn.close()
    if row:
        return jsonify({"question_text": row["question_text"]})
    return jsonify({"error":"Question not found"}), 404

# ---------- Chatbot Routes ----------

@app.route("/chatbot")
def chatbot_redirect():
    if "user" not in session:
        return redirect(url_for("login"))

    conn = get_db_connection()
    cur = conn.cursor(dictionary=True)
    cur.execute(
        "SELECT id FROM chat_sessions WHERE user_id=%s ORDER BY updated_at DESC LIMIT 1",
        (session["user_id"],)
    )
    row = cur.fetchone()
    cur.close()
    conn.close()

    if row:
        return redirect(url_for("chatbot_session", chat_id=row["id"]))
    else:
        return redirect(url_for("chatbot_new"))

@app.route("/chatbot/new", methods=["POST", "GET"])
def chatbot_new():
    if "user" not in session:
        return redirect(url_for("login"))

    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute(
        "INSERT INTO chat_sessions (user_id, title, created_at, updated_at) VALUES (%s, %s, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)",
        (session["user_id"], "New Chat")
    )
    chat_id = cur.lastrowid
    conn.commit()
    cur.close()
    conn.close()
    return redirect(url_for("chatbot_session", chat_id=chat_id))

@app.route("/chatbot/<int:chat_id>", methods=["GET", "POST"])
def chatbot_session(chat_id):
    if "user" not in session:
        return redirect(url_for("login"))

    conn = get_db_connection()
    cur = conn.cursor(dictionary=True)
    cur.execute(
        "SELECT * FROM chat_sessions WHERE id=%s AND user_id=%s",
        (chat_id, session["user_id"])
    )
    chat = cur.fetchone()
    if not chat:
        cur.close()
        conn.close()
        return redirect(url_for("chatbot_redirect"))

    cur.execute(
        "SELECT id, title FROM chat_sessions WHERE user_id=%s ORDER BY updated_at DESC",
        (session["user_id"],)
    )
    sessions = cur.fetchall()

    cur.execute("SELECT role, content FROM chat_messages WHERE session_id=%s ORDER BY created_at ASC", (chat_id,))
    messages = cur.fetchall()

    # Handle new message
    if request.method == "POST":
        user_msg = request.form.get("message", "").strip()
        if user_msg:
            cur.execute(
                "INSERT INTO chat_messages (session_id, role, content, created_at) VALUES (%s,%s,%s,CURRENT_TIMESTAMP)",
                (chat_id, "user", user_msg)
            )

            # Auto-rename chat if still "New Chat"
            if chat["title"] == "New Chat":
                cur.execute(
                    "UPDATE chat_sessions SET title=%s WHERE id=%s",
                    (user_msg[:30] + "...", chat_id)
                )

            # Prepare for AI response
            reply = "AI service unavailable."
            if client:
                try:
                    system_prompt = (
                        "You are an expert interview coach. "
                        "Provide clear, structured, and actionable advice with examples."
                    )
                    history = [{"role": "system", "content": system_prompt}]
                    for m in messages:
                        history.append({"role": m["role"], "content": m["content"]})
                    history.append({"role": "user", "content": user_msg})

                    response = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=history
                    )
                    reply = response.choices[0].message.content
                except Exception as e:
                    print("⚠️ AI error:", e)

            # Insert AI reply
            cur.execute(
                "INSERT INTO chat_messages (session_id, role, content, created_at) VALUES (%s,%s,%s,CURRENT_TIMESTAMP)",
                (chat_id, "assistant", reply)
            )

            # Update session timestamp
            cur.execute(
                "UPDATE chat_sessions SET updated_at=CURRENT_TIMESTAMP WHERE id=%s", (chat_id,)
            )
            conn.commit()
            cur.close()
            conn.close()
            return redirect(url_for("chatbot_session", chat_id=chat_id))

    cur.close()
    conn.close()
    return render_template(
        "chatbot.html",
        sessions=sessions,
        messages=messages,
        current_session_id=chat_id
    )

@app.route("/chatbot/delete/<int:chat_id>", methods=["POST"])
def chatbot_delete(chat_id):
    if "user" not in session:
        return redirect(url_for("login"))

    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute("DELETE FROM chat_messages WHERE session_id=%s", (chat_id,))
    cur.execute("DELETE FROM chat_sessions WHERE id=%s AND user_id=%s", (chat_id, session["user_id"]))
    conn.commit()
    cur.close()
    conn.close()
    return redirect(url_for("chatbot_redirect"))

@app.route('/logout')
def logout():
    session.clear()
    flash("Logged out successfully!", "info")
    return redirect(url_for('login'))

# Enhanced API endpoints
@app.route('/api/user_progress')
def api_user_progress():
    if 'user' not in session:
        return jsonify({"error": "Unauthorized"}), 403
    
    user_id = session.get('user_id')
    conn = get_db_connection()
    if not conn:
        return jsonify({"error": "Database error"}), 500
    
    try:
        cur = conn.cursor(dictionary=True)
        
        # Get comprehensive progress data
        cur.execute("""
            SELECT 
                AVG(score_overall) as avg_score,
                AVG(eye_contact_percent) as avg_eye_contact,
                AVG(confidence_score) as avg_confidence,
                COUNT(*) as total_attempts,
                DATE(MAX(attempted_at)) as last_attempt
            FROM live_interview_attempts 
            WHERE user_id = %s
        """, (user_id,))
        
        progress = cur.fetchone()
        return jsonify(progress or {})
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if conn:
            conn.close()

@app.route('/api/personalized_tips')
def api_personalized_tips():
    if 'user' not in session:
        return jsonify({"error": "Unauthorized"}), 403
    
    user_id = session.get('user_id')
    user_role = session.get('role')
    
    conn = get_db_connection()
    if not conn:
        return jsonify({"error": "Database error"}), 500
    
    try:
        cur = conn.cursor(dictionary=True)
        
        # Analyze user's weak areas
        cur.execute("""
            SELECT 
                AVG(eye_contact_percent) as avg_eye,
                AVG(confidence_score) as avg_confidence,
                AVG(score_overall) as avg_score,
                AVG(filler_word_count) as avg_fillers
            FROM live_interview_attempts 
            WHERE user_id = %s AND attempted_at >= DATE_SUB(NOW(), INTERVAL 30 DAY)
        """, (user_id,))
        
        stats = cur.fetchone()
        tips = []
        
        if stats:
            if stats['avg_eye'] < 60:
                tips.append("Practice maintaining eye contact by looking directly at the camera")
            if stats['avg_confidence'] < 0.6:
                tips.append("Work on speaking more clearly and at a steady pace")
            if stats['avg_fillers'] > 5:
                tips.append("Practice reducing filler words like 'um', 'uh', and 'like'")
            if stats['avg_score'] < 6:
                tips.append(f"Focus on studying {user_role}-specific technical concepts")
        
        return jsonify({"tips": tips})
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if conn:
            conn.close()
if __name__ == '__main__':
    app.run(debug=True)
